import asyncio
import subprocess
import tempfile
from pathlib import Path

from docx import Document

from email_agent.document.host import HostDocumentTools
from email_agent.sandbox.inmemory_environment import InMemoryEnvironment


async def test_pandoc_stages_inputs_and_copies_declared_outputs() -> None:
    env = InMemoryEnvironment()
    await env.write_text("docs/in.md", "# Hello\n")
    calls: list[list[str]] = []

    def runner(command, *, cwd, capture_output, text, timeout, check):
        calls.append(command)
        (cwd / "docs" / "out.html").write_text("<h1>Hello</h1>")
        return subprocess.CompletedProcess(command, 0, stdout="ok", stderr="")

    tools = HostDocumentTools(pandoc_path="pandoc", runner=runner)

    result = await tools.pandoc(
        env,
        args=["/workspace/docs/in.md", "-o", "/workspace/docs/out.html"],
        input_paths=["docs/in.md"],
        output_paths=["docs/out.html"],
    )

    assert "exit_code=0" in result
    assert calls[0][0] == "pandoc"
    assert calls[0][1].endswith("/docs/in.md")
    assert calls[0][3].endswith("/docs/out.html")
    assert await env.read_text("docs/out.html") == "<h1>Hello</h1>"


async def test_python_docx_sets_margins_and_writes_output() -> None:
    env = InMemoryEnvironment()
    with tempfile.TemporaryDirectory() as tmp:
        source_path = f"{tmp}/source.docx"
        output_path = f"{tmp}/output.docx"
        source = Document()
        source.add_paragraph("Hello")
        source.save(source_path)
        await env.write_bytes("docs/in.docx", await asyncio.to_thread(Path(source_path).read_bytes))

        tools = HostDocumentTools()

        result = await tools.python_docx(
            env,
            path="docs/in.docx",
            operations=[{"action": "set_margins", "all": 0.7}],
            output_path="docs/out.docx",
        )

        assert result == "wrote /workspace/docs/out.docx"
        await asyncio.to_thread(
            Path(output_path).write_bytes, await env.read_bytes("docs/out.docx")
        )
        output = Document(output_path)
        left_margin = output.sections[0].left_margin
        assert left_margin is not None
        assert round(left_margin.inches, 1) == 0.7
        assert output.paragraphs[0].text == "Hello"
