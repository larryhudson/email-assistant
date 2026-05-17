import asyncio
import subprocess
import tempfile
from collections.abc import Callable
from pathlib import Path, PurePosixPath

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.shared import Inches

from email_agent.sandbox.environment import SandboxEnvironment

WORKSPACE_ROOT = "/workspace"


class HostDocumentTools:
    """Host-side document tools for sandboxes that cannot run office binaries."""

    def __init__(
        self,
        *,
        pandoc_path: str = "pandoc",
        soffice_path: str = "soffice",
        timeout_seconds: int = 120,
        runner: Callable[..., subprocess.CompletedProcess[str]] | None = None,
    ) -> None:
        self._pandoc_path = pandoc_path
        self._soffice_path = soffice_path
        self._timeout_seconds = timeout_seconds
        self._runner = runner or subprocess.run

    async def pandoc(
        self,
        env: SandboxEnvironment,
        *,
        args: list[str],
        input_paths: list[str],
        output_paths: list[str],
        timeout_s: int | None = None,
    ) -> str:
        return await self._run_cli(
            env,
            binary=self._pandoc_path,
            args=args,
            input_paths=input_paths,
            output_paths=output_paths,
            timeout_s=timeout_s,
        )

    async def soffice(
        self,
        env: SandboxEnvironment,
        *,
        args: list[str],
        input_paths: list[str],
        output_paths: list[str],
        timeout_s: int | None = None,
    ) -> str:
        return await self._run_cli(
            env,
            binary=self._soffice_path,
            args=args,
            input_paths=input_paths,
            output_paths=output_paths,
            timeout_s=timeout_s,
        )

    async def python_docx(
        self,
        env: SandboxEnvironment,
        *,
        path: str,
        operations: list[dict],
        output_path: str | None = None,
    ) -> str:
        input_workspace_path = _workspace_path(path)
        output_workspace_path = _workspace_path(output_path or path)
        with tempfile.TemporaryDirectory(prefix="email-agent-docx-") as tmp:
            root = Path(tmp) / "workspace"
            input_host_path = _host_path(root, input_workspace_path)
            output_host_path = _host_path(root, output_workspace_path)
            input_host_path.parent.mkdir(parents=True, exist_ok=True)
            input_host_path.write_bytes(await env.read_bytes(input_workspace_path))

            await asyncio.to_thread(
                _apply_docx_operations,
                input_host_path,
                output_host_path,
                operations,
            )
            await env.write_bytes(output_workspace_path, output_host_path.read_bytes())

        return f"wrote {output_workspace_path}"

    async def _run_cli(
        self,
        env: SandboxEnvironment,
        *,
        binary: str,
        args: list[str],
        input_paths: list[str],
        output_paths: list[str],
        timeout_s: int | None,
    ) -> str:
        if not output_paths:
            return "ERROR: document tool failed\noutput_paths must include at least one file"

        with tempfile.TemporaryDirectory(prefix="email-agent-doc-") as tmp:
            root = Path(tmp) / "workspace"
            root.mkdir(parents=True)
            for path in input_paths:
                workspace_path = _workspace_path(path)
                host_path = _host_path(root, workspace_path)
                host_path.parent.mkdir(parents=True, exist_ok=True)
                host_path.write_bytes(await env.read_bytes(workspace_path))

            command = [binary, *[_host_arg(root, arg) for arg in args]]
            result = await asyncio.to_thread(
                self._runner,
                command,
                cwd=root,
                capture_output=True,
                text=True,
                timeout=timeout_s or self._timeout_seconds,
                check=False,
            )

            copied: list[str] = []
            for path in output_paths:
                workspace_path = _workspace_path(path)
                host_path = _host_path(root, workspace_path)
                if not host_path.exists():
                    continue
                await env.write_bytes(workspace_path, host_path.read_bytes())
                copied.append(workspace_path)

        stdout = str(result.stdout or "").strip()
        stderr = str(result.stderr or "").strip()
        lines = [f"exit_code={result.returncode}"]
        if copied:
            lines.append("outputs:")
            lines.extend(f"- {path}" for path in copied)
        else:
            lines.append("outputs: <none copied>")
        lines.extend(["stdout:", stdout, "stderr:", stderr])
        return "\n".join(lines)


def _apply_docx_operations(input_path: Path, output_path: Path, operations: list[dict]) -> None:
    doc = Document(str(input_path))
    for operation in operations:
        action = operation.get("action")
        if action == "set_margins":
            _set_margins(doc, operation)
        elif action == "set_orientation":
            _set_orientation(doc, operation)
        elif action == "replace_text":
            _replace_text(doc, str(operation["old"]), str(operation["new"]))
        else:
            raise ValueError(f"unsupported python_docx action: {action!r}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _set_margins(doc: DocxDocument, operation: dict) -> None:
    sections = doc.sections
    for section in sections:
        all_margin = operation.get("all")
        top = operation.get("top", all_margin)
        right = operation.get("right", all_margin)
        bottom = operation.get("bottom", all_margin)
        left = operation.get("left", all_margin)
        if top is not None:
            section.top_margin = Inches(float(top))
        if right is not None:
            section.right_margin = Inches(float(right))
        if bottom is not None:
            section.bottom_margin = Inches(float(bottom))
        if left is not None:
            section.left_margin = Inches(float(left))


def _set_orientation(doc: DocxDocument, operation: dict) -> None:
    value = str(operation.get("value", "")).lower()
    if value not in {"portrait", "landscape"}:
        raise ValueError("set_orientation value must be 'portrait' or 'landscape'")
    for section in doc.sections:
        if value == "landscape" and section.orientation != WD_ORIENT.LANDSCAPE:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        elif value == "portrait" and section.orientation != WD_ORIENT.PORTRAIT:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width, section.page_height = section.page_height, section.page_width


def _replace_text(doc: DocxDocument, old: str, new: str) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)


def _host_arg(root: Path, arg: str) -> str:
    if arg.startswith(f"{WORKSPACE_ROOT}/") or arg == WORKSPACE_ROOT:
        return str(_host_path(root, arg))
    return arg.replace(f"{WORKSPACE_ROOT}/", f"{root}/")


def _host_path(root: Path, workspace_path: str) -> Path:
    rel = _relative_workspace_path(workspace_path)
    return root / rel


def _workspace_path(path: str) -> str:
    if path.startswith(WORKSPACE_ROOT):
        return str(PurePosixPath(path))
    if path.startswith("/"):
        return str(PurePosixPath(f"{WORKSPACE_ROOT}{path}"))
    return str(PurePosixPath(WORKSPACE_ROOT) / PurePosixPath(path))


def _relative_workspace_path(path: str) -> PurePosixPath:
    workspace_path = PurePosixPath(_workspace_path(path))
    rel = workspace_path.relative_to(WORKSPACE_ROOT)
    if ".." in rel.parts:
        raise ValueError(f"invalid workspace path: {path}")
    return rel
