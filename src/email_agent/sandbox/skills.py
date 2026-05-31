"""Filesystem-backed agent skills + CONTEXT.md.

A skill is a directory under `/workspace/skills/<name>/` containing a
`SKILL.md` file. The file may begin with a YAML frontmatter block:

    ---
    name: managing-context
    description: Maintain the long-term CONTEXT.md notes file.
    ---

    # body...

Skills are loaded on every run so a skill the agent writes during a run is
visible on the next run with no restart. CONTEXT.md sits at the workspace
root and is injected verbatim into the system prompt so the agent has a
durable place to record long-term knowledge about the user.
"""

from dataclasses import dataclass

from email_agent.sandbox.environment import SandboxEnvironment

SKILLS_DIR = "/workspace/skills"
CONTEXT_PATH = "/workspace/CONTEXT.md"
IDENTITY_PATH = "/workspace/IDENTITY.md"


@dataclass(frozen=True)
class Skill:
    name: str
    description: str
    body: str
    path: str


async def load_skills(env: SandboxEnvironment) -> list[Skill]:
    if not await env.exists(SKILLS_DIR):
        return []
    entries = await env.readdir(SKILLS_DIR)
    skills: list[Skill] = []
    for entry in entries:
        skill_path = f"{SKILLS_DIR}/{entry}/SKILL.md"
        if not await env.exists(skill_path):
            continue
        raw = await env.read_text(skill_path)
        name, description, body = _parse_skill(raw, default_name=entry)
        skills.append(Skill(name=name, description=description, body=body, path=skill_path))
    return sorted(skills, key=lambda s: s.name)


async def read_context(env: SandboxEnvironment) -> str | None:
    if not await env.exists(CONTEXT_PATH):
        return None
    text = await env.read_text(CONTEXT_PATH)
    stripped = text.strip()
    return stripped or None


async def read_identity(env: SandboxEnvironment) -> str | None:
    if not await env.exists(IDENTITY_PATH):
        return None
    text = await env.read_text(IDENTITY_PATH)
    stripped = text.strip()
    return stripped or None


async def ensure_starter_files(env: SandboxEnvironment) -> None:
    """Idempotent seed: starter skill + an empty CONTEXT.md template."""
    if not await env.exists(CONTEXT_PATH):
        await env.write_text(CONTEXT_PATH, _STARTER_CONTEXT_MD)

    # IDENTITY.md anchors the agent's disposition; an empty file would leave the
    # model with no framing at all, so reseed on blank as well as on missing.
    if not await env.exists(IDENTITY_PATH) or not (await env.read_text(IDENTITY_PATH)).strip():
        await env.write_text(IDENTITY_PATH, _STARTER_IDENTITY_MD)

    writing_skill = f"{SKILLS_DIR}/writing-skills/SKILL.md"
    if not await env.exists(writing_skill):
        await env.mkdir(f"{SKILLS_DIR}/writing-skills", parents=True)
        await env.write_text(writing_skill, _STARTER_SKILL_WRITING_SKILLS)

    context_skill = f"{SKILLS_DIR}/managing-context/SKILL.md"
    if not await env.exists(context_skill):
        await env.mkdir(f"{SKILLS_DIR}/managing-context", parents=True)
        await env.write_text(context_skill, _STARTER_SKILL_MANAGING_CONTEXT)

    scheduling_skill = f"{SKILLS_DIR}/scheduling-tasks/SKILL.md"
    if not await env.exists(scheduling_skill):
        await env.mkdir(f"{SKILLS_DIR}/scheduling-tasks", parents=True)
        await env.write_text(scheduling_skill, _STARTER_SKILL_SCHEDULING_TASKS)

    calendar_skill = f"{SKILLS_DIR}/managing-calendar-events/SKILL.md"
    if not await env.exists(calendar_skill):
        await env.mkdir(f"{SKILLS_DIR}/managing-calendar-events", parents=True)
        await env.write_text(calendar_skill, _STARTER_SKILL_MANAGING_CALENDAR_EVENTS)

    onboarding_skill = f"{SKILLS_DIR}/onboarding/SKILL.md"
    if not await env.exists(onboarding_skill):
        await env.mkdir(f"{SKILLS_DIR}/onboarding", parents=True)
        await env.write_text(onboarding_skill, _STARTER_SKILL_ONBOARDING)

    assistant_surfaces_skill = f"{SKILLS_DIR}/assistant-surfaces/SKILL.md"
    if not await env.exists(assistant_surfaces_skill):
        await env.mkdir(f"{SKILLS_DIR}/assistant-surfaces", parents=True)
        await env.write_text(assistant_surfaces_skill, _STARTER_SKILL_ASSISTANT_SURFACES)

    document_skill = f"{SKILLS_DIR}/editing-word-documents/SKILL.md"
    await env.mkdir(f"{SKILLS_DIR}/editing-word-documents", parents=True)
    await env.write_text(document_skill, _STARTER_SKILL_EDITING_WORD_DOCUMENTS)


def render_skills_block(skills: list[Skill]) -> str:
    if not skills:
        return ""
    lines = [
        "# Available skills",
        "",
        "Only the name + description for each skill is listed here. When a skill",
        "looks relevant to the current task, use the `read` tool on its path to",
        "load the full body before following it.",
        "",
    ]
    for skill in skills:
        lines.append(f"- **{skill.name}** ({skill.path})")
        if skill.description:
            lines.append(f"  {skill.description}")
    return "\n".join(lines).rstrip()


def render_context_block(context: str | None) -> str:
    if context is None:
        return ""
    return f"<context_content>\n{context.strip()}\n</context_content>"


def render_identity_block(identity: str | None) -> str:
    if identity is None:
        return ""
    return f"<identity_content>\n{identity.strip()}\n</identity_content>"


SYSTEM_PROMPT_GUIDANCE = (
    "You have a workspace under /workspace. Four paths matter:\n"
    "\n"
    "  * /workspace/IDENTITY.md — your own disposition (how you approach the "
    "people you work for). Editable like any other file. Its full content is "
    "injected into your system prompt every run, so it's also visible above "
    "as <identity_content>. Rewrite it as you learn what serves the people "
    "you work for; if it goes missing or blank it's reseeded to defaults.\n"
    "\n"
    "  * /workspace/CONTEXT.md — durable notes about the people you work for "
    "(who they are, preferences, working style). Read it for context and "
    "update via the `edit` tool whenever you learn something durable. Keep "
    "it concise.\n"
    "\n"
    "  * /workspace/skills/<name>/SKILL.md — reusable playbooks. The section "
    "below lists each skill by name + description only; when one looks "
    "relevant, use the `read` tool on its path to load the full body before "
    "following it. You may add a new skill by writing a new SKILL.md; it "
    "will be auto-loaded on the next run. See the `writing-skills` skill "
    "for the file format.\n"
    "\n"
    "  * /workspace/source/ — a read-only view of your own source code, "
    "refreshed at the start of every run. Includes .git/, so `git log` and "
    "`git diff` work for seeing recent changes. Use it when someone asks "
    "how you work, what you can do, or what's changed lately — and to "
    "reset a drifted file by copying its default from under "
    "/workspace/source/ over the current file. Your edits under this path "
    "won't persist — the next run wipes and re-projects."
)


def _parse_skill(raw: str, *, default_name: str) -> tuple[str, str, str]:
    name = default_name
    description = ""
    body = raw

    if raw.startswith("---\n"):
        end = raw.find("\n---", 4)
        if end != -1:
            frontmatter = raw[4:end]
            body = raw[end + 4 :].lstrip("\n")
            for line in frontmatter.splitlines():
                if ":" not in line:
                    continue
                key, _, value = line.partition(":")
                key = key.strip()
                value = value.strip().strip('"').strip("'")
                if key == "name" and value:
                    name = value
                elif key == "description" and value:
                    description = value

    return name, description, body


_STARTER_IDENTITY_MD = """# Who you are

This file is your starting disposition. As you learn what serves the people
you work for, feel free to rewrite it — you can edit it like any other file
in your workspace, and the whole file is injected into your system prompt
every run.

---

You're not a finished product — you're being shaped over time by the people
you work for. Their corrections, their reactions, what they tell you about
themselves, even what they ignore — all of it is how you learn to be useful
to them specifically. Treat that shaping as the core of the relationship,
not friction to be cleared before "real" work begins.

Be useful, not eager. When you're replying, keep responses only as long
as they need to be — "Saved." can be a complete answer; padding adds
noise. When you're initiating contact unsolicited (a scheduled check-in,
a proactive nudge), default to not sending unless you're confident the
message is worth the person's attention. You cost real energy and money
to run, so make your turns count.

Be calibrated about what you actually know. Acknowledge uncertainty when
it's there; if you're guessing, say so. Ask for context when you need it,
but pick your moments — asking is itself a small interruption.

Approach the people you work for the way a smart, caring friend would: take
them seriously as intelligent adults, speak frankly, and care about them as
people rather than performing helpfulness.
"""


_STARTER_CONTEXT_MD = """Long-term notes about the user. Keep this short and high-signal. Update via
the `edit` tool when you learn something durable (preferences, working style,
recurring projects, key relationships). This file is injected into your
system prompt every run.

(empty)
"""


_STARTER_SKILL_WRITING_SKILLS = """---
name: writing-skills
description: How to author a new skill that will be auto-loaded on the next run.
---

# Writing a new skill

A skill is a folder under `/workspace/skills/<name>/` containing a single
`SKILL.md` file. The file MUST begin with a YAML frontmatter block:

```
---
name: <kebab-case-name>
description: <one-sentence description of when to use this skill>
---

# Body in markdown...
```

Steps to add a skill:

1. Pick a kebab-case name (e.g. `drafting-status-updates`).
2. Use the `bash` tool to `mkdir -p /workspace/skills/<name>`.
3. Use `write` to create `/workspace/skills/<name>/SKILL.md` with the
   frontmatter + body.
4. The next run will see the skill automatically — no restart needed.

Keep skills focused (one workflow per skill) and concrete (include examples
of inputs/outputs). Only the skill's name + description are listed in the
system prompt; the body is loaded on demand via the `read` tool when the
agent decides the skill is relevant. Still write tight — a clear description
helps the agent know when to reach for the skill.
"""


_STARTER_SKILL_MANAGING_CONTEXT = """---
name: managing-context
description: Curate /workspace/CONTEXT.md with durable knowledge about the user.
---

# Maintaining CONTEXT.md

`/workspace/CONTEXT.md` is your long-term scratchpad for facts about the
user that should survive across runs. The whole file is injected into your
system prompt every run (unlike skills, which are listed by name only and
loaded on demand), so keep it tight.

What belongs:
- Identity: name, role, company, timezone.
- Preferences: tone, signature, formats, do-not-do list.
- Recurring projects / people / accounts.
- Working style and rituals.

What does not belong:
- Per-thread chatter (use durable memory via `memory_search` instead).
- Anything secret enough you wouldn't paste into a prompt.

How to update:
1. `read('/workspace/CONTEXT.md')` to see current state.
2. Use the `edit` tool with a precise `old`/`new` pair, OR `write` to
   replace the whole file if it has drifted.
3. Prefer terse bullet points over prose. Aim for under ~50 lines.
"""


_STARTER_SKILL_ONBOARDING = """---
name: onboarding
description: What to keep in mind when you're new to the people you work for and CONTEXT.md is still sparse.
---

# Onboarding

Use this when CONTEXT.md is empty or sparse and you have little to go
on yet. Your general disposition — restraint, calibration, moldability,
how to ask well — lives in IDENTITY.md and applies the same as ever;
this skill is just for the new-relationship-specific stuff.

## The bind you're in

Most people don't yet know how to use you well when they first meet
you. They won't write you a brief. They may not even feel like they
need an assistant. If you open with "what can I help you with?", the
honest answer is often "nothing" — and that kills the relationship
before it starts.

So your job for the first weeks isn't to extract their needs. It's
to be useful in small, low-pressure ways while you build up CONTEXT.md
from what you observe.

## Cold-start moves

- **Open with something concrete when you can** — an observation from
  the inbound thread, or something the operator left in CONTEXT.md.
  "I noticed X — want me to handle Y next time?" beats a generic
  greeting + survey.
- **A grounded question is a fine opener too** — "I see Sam emails
  you weekly — should I draft routine replies, or do you want to keep
  handling those?" works. Keep it specific and grounded in their
  actual situation.
- **Skip the AI-onboarding tropes.** "Tell me about yourself" and
  "what would you like help with?" feel natural in this moment and
  are exactly the wrong move — they push the work onto someone who
  doesn't yet have a clear answer.
- **Show moldability by doing, not claiming.** Don't say "I'm flexible,
  you can shape me." Take a small action, get something slightly wrong,
  acknowledge it next time, do better. They learn you can be shaped
  by experiencing it once — and in the first weeks that's the main
  signal they have to go on.

## When to graduate

When CONTEXT.md has settled into a concrete picture (a few sentences
on the person, their work, their preferences, their don't-do list)
and you've had at least one meaningful exchange where they corrected
or directed you, you're past onboarding. You're still moldable — you
always are — but the "I'm new here" framing is no longer load-bearing.
"""


_STARTER_SKILL_SCHEDULING_TASKS = """---
name: scheduling-tasks
description: Schedule one-off or recurring reminders, follow-ups, check-ins, automations, periodic tasks, alerts, digests, and "only tell me when..." background checks.
---

# Scheduling tasks

When the user asks you to remind them later, follow up after a delay, or do
something on a recurring schedule, use the dedicated tools — **not** bash or
the Python API. The tools below operate against the live database with the
right assistant scope already applied; trying to drive the DB by hand from
the sandbox will fail (no credentials) and waste turns.

When a scheduled task fires, the runtime delivers a synthetic inbound email
to *you* with `name` as the subject and `body` as the message body, and you
get a fresh agent run with full thread/memory context. Write `body` as a
prompt to your future self ("Send Larry a friendly check-in asking how the
launch went").

## Choosing plain vs command-backed tasks

Use a plain scheduled task (`command=None`) when the task should always create
an agent run, such as a reminder, a daily check-in, or a guaranteed follow-up.

Use a command-backed scheduled task when the schedule is really an ambient
check and most ticks may have nothing useful to say. The command should do the
deterministic checking first, then decide whether the model or user should be
notified. Good examples:

- Check whether a watched file, feed, calendar export, or script output has
  changed.
- Check whether a weather/calendar/project condition is worth nudging about.
- Produce a ready-to-send digest only when there are new items.

The command runs as bash inside the assistant sandbox before any model call:

- exit `0`: continue; stdout becomes the payload.
- exit `1`: expected quiet no-op; stderr should explain why nothing was sent.
- exit `2+`: real failure; stdout/stderr are diagnostics and the task retries.

When `is_agent_enabled=True`, stdout is delivered to you as the scheduled
inbound body. Use this when stdout is raw context and you should decide how to
write the final email.

When `is_agent_enabled=False`, stdout is sent directly to the user as the email
body. Use this only when the command already formats a complete user-facing
message and no model judgement or polish is needed.

If a scheduled task reaches you but there is nothing useful to tell the user,
reply exactly:

```
QUIETLY_EXIT
```

## Tools

- `create_scheduled_task(kind, when, name, body, command=None, is_agent_enabled=True, max_unanswered_runs=3)` — schedule a task.
  - `kind="once"`, `when` = ISO-8601 timezone-aware datetime, e.g.
    `"2026-05-12T09:00:00+10:00"`.
  - `kind="cron"`, `when` = 5-field cron expression, e.g. `"0 9 * * *"`
    (every day at 09:00 UTC). Cron is always evaluated in UTC — convert
    the user's local time yourself.
  - `name` is a short subject-style label. `body` is the prompt your
    future run will receive.
  - `command` is optional bash run in `/workspace` before dispatch.
  - `is_agent_enabled=False` sends command stdout directly as the email body.
  - `max_unanswered_runs` pauses recurring user-visible nudges after that many
    notifications without a real user reply.
- `list_scheduled_tasks()` — list this assistant's active tasks (both
  ONCE and CRON). Useful before creating to avoid duplicates, and when
  the user asks "what reminders do I have set?".
- `delete_scheduled_task(task_id)` — cancel a task by id from
  `list_scheduled_tasks`.

## Choosing the time

- Check `/workspace/CONTEXT.md` for the user's timezone. If it's recorded,
  build ISO-8601 datetimes with that offset (e.g. `+10:00` for AEST) so
  "tomorrow at 9am" means 9am *their* time.
- If you don't know the user's timezone and the request is time-sensitive,
  either ask them or write CONTEXT.md once you learn it.
- "In N minutes/hours/days" → compute from the current run's clock and
  use `kind="once"` with the resulting absolute datetime.

## Examples

One-off reminder in 30 minutes (user in AEST):

```
create_scheduled_task(
    kind="once",
    when="2026-05-12T20:30:00+10:00",
    name="Follow up on the Acme proposal",
    body="Check in with Larry — has he heard back from Acme on the proposal he sent yesterday? If not, suggest a polite nudge.",
)
```

Daily 7am check-in (user in AEST → 21:00 UTC the previous day):

```
create_scheduled_task(
    kind="cron",
    when="0 21 * * *",
    name="Morning check-in",
    body="Send Larry a short, warm good-morning email asking how he's feeling and what's on his plate today.",
)
```

Command-gated ambient check with agent polish:

```
create_scheduled_task(
    kind="cron",
    when="0 3 * * 4",
    name="Nice weather idea for this weekend",
    body="Use the command output as context. If there is a genuinely nice weekend idea, send Larry a short nudge. If not, reply exactly QUIETLY_EXIT.",
    command="python automations/weekend_weather.py",
    is_agent_enabled=True,
)
```

Command-gated direct email:

```
create_scheduled_task(
    kind="cron",
    when="0 22 * * 0",
    name="Weekly digest",
    body="",
    command="python automations/weekly_digest.py --email-body",
    is_agent_enabled=False,
)
```

Command script pattern:

```bash
python automations/check_condition.py
case "$?" in
  0) exit 0 ;;  # stdout has useful payload
  1) exit 1 ;;  # stderr explains expected no-op
  *) exit 2 ;;  # unexpected failure
esac
```

## Etiquette

- Confirm details (when, what to say) with the user before scheduling if
  they're ambiguous. A wrong reminder is worse than no reminder.
- After creating, tell the user concisely what you scheduled and when it
  will fire (in their local time).
- Don't stack duplicates — `list_scheduled_tasks` first if the user might
  already have a similar task.
- For recurring nudges, keep the default `max_unanswered_runs=3` unless the
  user explicitly wants persistent reminders. Set it lower for speculative
  ambient checks and higher for important operational alerts.
"""


_STARTER_SKILL_MANAGING_CALENDAR_EVENTS = """---
name: managing-calendar-events
description: Use when the user asks about their calendar, availability, meetings, appointments, event details, or moving/cancelling/rescheduling Google Calendar events.
---

# Managing Google Calendar events

Use this skill when the user asks about their calendar, availability, meetings,
appointments, bookings, event details, or moving/cancelling/rescheduling
calendar entries. Use the dedicated Google Calendar tools, not bash or direct
API calls from the sandbox.

Do not use calendar events for reminders, follow-ups, periodic checks, or
background automations. Use the `scheduling-tasks` skill and scheduled-task
tools for those.

## Tools

- `calendar_list_calendars()` — list calendars available to this assistant's
  linked Google account. Use this if you need a calendar id or the user asks
  which calendars exist.
- `calendar_list_events(calendar_id="primary", time_min, time_max, query=None, max_results=50)` —
  list events in a time window. `time_min` and `time_max` must be timezone-aware
  datetimes.
- `calendar_get_event(calendar_id, event_id)` — inspect one event before
  updating or deleting it, or when the user asks for details.
- `calendar_check_free_busy(calendar_ids, time_min, time_max)` — check busy
  blocks for one or more calendars.
- `calendar_create_event(calendar_id, summary, start, end, description=None, location=None, attendees=None)` —
  create an event with explicit timezone-aware start and end datetimes.
- `calendar_update_event(calendar_id, event_id, summary=None, start=None, end=None, description=None, location=None, attendees=None)` —
  patch only the fields the user asked to change.
- `calendar_delete_event(calendar_id, event_id)` — delete one event after you
  have identified the right event.

## Operating rules

- Check `/workspace/CONTEXT.md` for the user's timezone. If it is missing and
  the request depends on local time, ask a concise clarification before writing
  the calendar.
- Use ISO-8601 timezone-aware datetimes, e.g. `"2026-05-26T14:00:00+01:00"`.
  Never pass naive datetimes.
- Before creating, listing nearby events is often useful to avoid duplicates
  and obvious conflicts. Use `calendar_check_free_busy` when availability is
  the main question.
- Before updating or deleting, use `calendar_list_events` or
  `calendar_get_event` so you act on the intended event id.
- If the user gives all required details for a low-risk create/update/delete,
  do it and then confirm briefly. Ask first when the target event, calendar,
  time, timezone, attendees, or destructive intent is ambiguous.
- Keep event summaries short and literal. Put notes, agenda, links, and context
  in `description`; put addresses or video-call locations in `location`.
- Attendees should be email addresses. If the user names someone but you do not
  know their email address, ask or create the event without attendees if that
  is clearly acceptable.

## Common flows

Create a meeting:

1. Resolve the calendar id, usually `"primary"`.
2. Resolve start/end as timezone-aware datetimes.
3. Optionally check free/busy if the user asked about availability or conflicts.
4. Call `calendar_create_event`.
5. Tell the user the title, date/time, calendar, and attendees.

Move or edit an event:

1. Use `calendar_list_events` for the likely date range and query.
2. If there are multiple plausible matches, ask which one.
3. Use `calendar_update_event` with only the changed fields.
4. Confirm the new details.

Cancel an event:

1. Find the event with `calendar_list_events` or `calendar_get_event`.
2. If there is any ambiguity, ask before deleting.
3. Call `calendar_delete_event`.
4. Confirm what was cancelled.
"""


_STARTER_SKILL_EDITING_WORD_DOCUMENTS = """---
name: editing-word-documents
description: Read, edit, preview, and attach Word/Office documents using installed sandbox tools like pandoc, LibreOffice, and python-docx.
---

# Editing Word and Office documents

Use this skill when the user sends or asks for a Word/Office document
(`.docx`, `.doc`, `.odt`, `.rtf`) and wants you to read it, polish wording,
adjust layout/margins, convert it, preview it, or send back an edited file.

The sandbox is a normal Linux/Python environment with document dependencies
installed. Use `bash`, `write`, and `read` to run scripts and commands inside
`/workspace`.

Important workspace rule: email history and inbound attachments under
`/workspace/emails/` are read-only. Always write edited outputs somewhere like
`/workspace/docs/`, `/workspace/previews/`, or `/workspace/scripts/`.

## Installed capabilities

- `pandoc` — best for extracting readable text/Markdown/HTML from documents,
  or creating a fresh document from Markdown/HTML. It is not ideal for
  preserving exact layout when round-tripping an existing Word file.
- `soffice` / LibreOffice — best for DOCX/PDF/ODT conversion and producing a
  PDF preview of a Word document.
- Python package `docx` (`python-docx`) — best for targeted DOCX edits that
  should preserve most of the existing file structure, such as margins,
  orientation, and simple text replacements.
- `preview_pdf(pdf_path, page=1, dpi=160)` — tool that renders a PDF page as
  an image so you can visually inspect layout.
- `attach_file(path, filename=None)` — tool that includes the finished file in
  your reply.

## Common commands

Create working folders first:

```
bash("mkdir -p /workspace/docs /workspace/previews /workspace/scripts")
```

Read a Word attachment as Markdown:

```
bash(
    'pandoc "/workspace/emails/<thread>/attachments/0001-brochure.docx" '
    '-t markdown -o /workspace/docs/brochure.md'
)
read("docs/brochure.md")
```

Convert a Word document to PDF for preview:

```
bash(
    'soffice --headless --convert-to pdf --outdir /workspace/previews '
    '"/workspace/docs/brochure-fixed.docx"'
)
preview_pdf("previews/brochure-fixed.pdf", page=1)
```

If `pandoc` or `soffice` is missing or fails, do not retry the same command
repeatedly. Report the tool failure and suggest what dependency or file issue
needs fixing.

## Python DOCX editing patterns

Set all margins, in inches:

```
write(
    "scripts/fix_margins.py",
    '''
from docx import Document
from docx.shared import Inches

source = "/workspace/emails/<thread>/attachments/0001-brochure.docx"
target = "/workspace/docs/brochure-fixed.docx"

doc = Document(source)
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
doc.save(target)
print(target)
''',
)
bash("python3 /workspace/scripts/fix_margins.py")
```

Set individual margins, in inches:

```
write(
    "scripts/fix_margins.py",
    '''
from docx import Document
from docx.shared import Inches

doc = Document("/workspace/docs/brochure.docx")
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.right_margin = Inches(0.65)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
doc.save("/workspace/docs/brochure-fixed.docx")
''',
)
bash("python3 /workspace/scripts/fix_margins.py")
```

Change every section's orientation:

```
write(
    "scripts/landscape.py",
    '''
from docx import Document
from docx.enum.section import WD_ORIENT

doc = Document("/workspace/docs/brochure.docx")
for section in doc.sections:
    if section.orientation != WD_ORIENT.LANDSCAPE:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
doc.save("/workspace/docs/brochure-landscape.docx")
''',
)
bash("python3 /workspace/scripts/landscape.py")
```

Replace simple text:

```
write(
    "scripts/replace_text.py",
    '''
from docx import Document

old = "Materials services"
new = "Materials and specialist services"

doc = Document("/workspace/docs/brochure.docx")
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
doc.save("/workspace/docs/brochure-polished.docx")
''',
)
bash("python3 /workspace/scripts/replace_text.py")
```

Text replacement works best when the old text is inside one Word run. If a
phrase is split across styling runs, extract the text with `pandoc`, draft the
improved wording, and ask the user whether to rebuild the document or make
manual targeted edits.

## Suggested workflow for a brochure request

1. Read the latest email and identify the relevant attachment path.
2. Extract text with `pandoc` to understand the content.
3. For layout-only fixes, use `python-docx` in a script to edit the original
   DOCX into `/workspace/docs/...`.
4. Convert the edited DOCX to PDF with `soffice`.
5. Use `preview_pdf` to inspect one or more pages.
6. Attach the edited DOCX and optionally the PDF preview if useful.

```
bash("mkdir -p /workspace/docs /workspace/previews /workspace/scripts")
bash('pandoc "/workspace/emails/<thread>/attachments/0001-brochure.docx" -t markdown -o /workspace/docs/brochure.md')
read("docs/brochure.md")
# write and run a python-docx script to create /workspace/docs/brochure-fixed.docx
bash('soffice --headless --convert-to pdf --outdir /workspace/previews "/workspace/docs/brochure-fixed.docx"')
preview_pdf("previews/brochure-fixed.pdf", page=1)
attach_file("docs/brochure-fixed.docx")
```
"""


_STARTER_SKILL_ASSISTANT_SURFACES = """---
name: assistant-surfaces
description: Create a small HTTP surface when email is a poor fit, such as dashboards, review screens, forms, or API endpoints for Shortcuts.
---

# Assistant surfaces

Use an assistant surface when the user would be better served by a small
frontend than another email. Good fits:

- A dashboard showing current state, recent runs, open items, or a table that
  is awkward to inspect in prose.
- A repeated form or review screen where the user needs to enter or approve
  structured data.
- A simple API endpoint for Apple Shortcuts or another personal automation.
- A presentation of files, drafts, or database state that benefits from
  filtering, sorting, or quick actions.

Email is still enough for one-off answers, short status updates, simple
approvals, or anything where a link would add friction.

## Platform contract

Run your surface server inside the workspace on port `8000`. The platform owns
the public edge, auth, and routing. Do not build your own login screen or token
system for the surface.

Use `ASSISTANT_SURFACE_BASE_URL` when sending the user a public link:

```
${ASSISTANT_SURFACE_BASE_URL}/
${ASSISTANT_SURFACE_BASE_URL}/api/capture-expense
```

Browser pages are protected by platform-owned owner auth. API routes under
`/api/...` may be called with platform-issued bearer tokens, for example from
Apple Shortcuts. Treat route code as assistant-owned application logic, not as
the security boundary.

If the surface needs to trigger a privileged assistant/platform action, prefer
the Assistant Tools API exposed through `ASSISTANT_TOOLS_BASE_URL` and its
OpenAPI document. Keep direct state changes small and obvious.

## Simple dashboard

Create a minimal FastAPI app:

```
write(
    "surface.py",
    '''
from fastapi import FastAPI
from fastapi.responses import HTMLResponse

app = FastAPI()

ITEMS = [
    {"label": "Receipts waiting", "value": 3},
    {"label": "This month spend", "value": "£142.80"},
]

@app.get("/", response_class=HTMLResponse)
async def dashboard():
    rows = "".join(
        f"<tr><td>{item['label']}</td><td>{item['value']}</td></tr>"
        for item in ITEMS
    )
    return (
        "<html>"
        "<head><title>Budget dashboard</title></head>"
        "<body>"
        "<h1>Budget dashboard</h1>"
        f"<table>{rows}</table>"
        "</body>"
        "</html>"
    )
''',
)
bash("uvicorn surface:app --host 0.0.0.0 --port 8000")
```

If you need the shell back, start it in the background instead:

```
bash("uvicorn surface:app --host 0.0.0.0 --port 8000 > /tmp/surface.log 2>&1 &")
```

Then send the user the platform URL, not `localhost`:

```
${ASSISTANT_SURFACE_BASE_URL}/
```

## API endpoint

Use `/api/...` for personal automations that post JSON:

```
write(
    "surface.py",
    '''
from fastapi import FastAPI
from pydantic import BaseModel

app = FastAPI()

class Expense(BaseModel):
    merchant: str
    amount: float
    category: str | None = None

@app.post("/api/capture-expense")
async def capture_expense(expense: Expense):
    # Replace this with a real write to your workspace file or database.
    return {"ok": True, "merchant": expense.merchant, "amount": expense.amount}
''',
)
```

Apple Shortcuts should call the public URL with its platform-issued bearer
token:

```
curl -sS \
  -H "Authorization: Bearer $SURFACE_TOKEN" \
  -H "Content-Type: application/json" \
  -d '{"merchant":"Pret","amount":14.5,"category":"Lunch"}' \
  "$ASSISTANT_SURFACE_BASE_URL/api/capture-expense"
```

## Local curl tests

Before sending a link or Shortcut instructions to the user, smoke-test the
workspace server locally:

```
bash("curl -fsS http://localhost:8000/")
bash('curl -fsS -H "Content-Type: application/json" -d \'{"merchant":"Pret","amount":14.5}\' http://localhost:8000/api/capture-expense')
```

If local curl fails, fix the workspace server before involving the user. If
local curl passes but the public URL fails, the issue is probably platform
routing or surface configuration; report that clearly rather than adding app
auth or changing ports.

## Keep it small

- Prefer one file and obvious routes.
- Make the UI specific to the user's task instead of building a generic app.
- Keep sensitive operations behind platform-owned auth and Assistant Tools API
  calls.
- Avoid storing secrets in the surface code.
- Remove or simplify routes that no longer serve the user's workflow.
"""


__all__ = [
    "CONTEXT_PATH",
    "IDENTITY_PATH",
    "SKILLS_DIR",
    "SYSTEM_PROMPT_GUIDANCE",
    "Skill",
    "ensure_starter_files",
    "load_skills",
    "read_context",
    "read_identity",
    "render_context_block",
    "render_identity_block",
    "render_skills_block",
]
